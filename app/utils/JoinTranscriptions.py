import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt

def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', s)]

def combine_text_files(input_folder, output_file, username):
    output_path = os.path.join(f"tmp_{username}", "joined_text")
    if not os.path.exists(output_path):
        os.makedirs(output_path)
    
    full_file_output = os.path.join(output_path, output_file)

    files = os.listdir(input_folder)

    text_files = [f for f in files if f.endswith(".txt")]
    text_files.sort(key=natural_sort_key)

    # Create or open the output file in write mode
    with open(full_file_output, 'w', encoding="utf-8") as outfile:

        # Iterate over all the files in the input folder
        for filename in text_files:
            # Only consider text files
            if filename.endswith('.txt'):
                # Construct full file path
                file_path = os.path.join(input_folder, filename)
                # Open each text file in read mode
                try:
                    # Open each text file in read mode
                    with open(file_path, "r", encoding="utf-8") as infile:
                        # Read the contents of the file and write it to the output file
                        outfile.write(infile.read())
                        # Add a newline character to separate the contents of different files
                        outfile.write("\n")
                except UnicodeDecodeError as e:
                    print(f"Skipping file {filename} due to encoding error: {e}")

def read_text_file(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def text_to_doc(text, output_path):
    doc = Document()

    styles = doc.styles
    style = styles['Normal']
    style.paragraph_format.space_after = Pt(0)
    style = styles['List Bullet']
    style.paragraph_format.space_after = Pt(0)

    # Split the text into lines
    lines = text.split('\n')

    # Loop through each line
    for line in lines:
        # Check if the line starts with '###'
        if line.startswith('###'):
            # Extract the header text (excluding '###')
            header_text = line.strip('# ')
            # Add the header to the document
            doc.add_heading(header_text, level=1).runs[0].font.size = Pt(14)
        elif line.startswith('  -'):
            # Add the line as a bullet point
            paragraph = doc.add_paragraph(style='List Bullet')
            add_bold_text(paragraph, line.strip('  -'))
        else:
            # Add the content as a paragraph
            paragraph = doc.add_paragraph()
            add_bold_text(paragraph, line)
    
    doc.save(output_path)

def add_bold_text(paragraph, text):
    # Regular expression to find text between **double asterisks**
    bold_texts = re.split(r'(\*\*.*?\*\*)', text)
    for part in bold_texts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(input_file_path, output_file_path):
    text = read_text_file(input_file_path)
    text_to_doc(text, output_file_path)

def summary_to_word_doc(input_file, username):

    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_folder = os.path.join(f"tmp_{username}", "word_summary")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    output_file_path = os.path.join(output_folder, f"{base_name}.docx")

    markdown_to_docx(input_file, output_file_path)

    return output_file_path


def json_to_word(input_file, username, json_data, title="Meeting Summary"):

    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_folder = os.path.join(f"tmp_{username}", "word_summary")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    output_file_path = os.path.join(output_folder, f"{base_name}.docx")

    # Create a new Document
    doc = Document()

    doc.add_heading(title, level=0)

    # Iterate over the JSON data
    for key, value in json_data.items():
        # Add section header
        doc.add_heading(key, level=1)

        if isinstance(value, str):
            # Add a paragraph for string values
            doc.add_paragraph(value)
        elif isinstance(value, list):
            # Add bullet points for list values
            for item in value:
                doc.add_paragraph(item, style='ListBullet')
    
    # Save the document
    doc.save(output_file_path)
    print(f"Document saved as {output_file_path}")

    return output_file_path
